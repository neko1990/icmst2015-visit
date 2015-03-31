"""
need "antiword"" cli-tool and "docx"" python lib to work
"""
import os
try:
    from docx import Document
except ImportError:
    class Document:
        def __init__(self,fn):
            """dummy class for docx"""
            self.fn = fn
            self.paragraphs = []


def doc_to_text_antiword(filename):
    (fi, fo, fe) = os.popen3('antiword -w 0 "%s"' % filename)
    fi.close()
    retval = fo.read()
    erroroutput = fe.read()
    fo.close()
    fe.close()
    if not erroroutput:
        return retval
    else:
        raise OSError("Executing the command caused an error: %s" % erroroutput)

def parse_user_name_aff(filename):
    try:
        txt = doc_to_text_antiword(filename).split('Abstract')[0]
    except OSError:
        print filename
        return ['DOC_ERROR','']
    pts=filter(lambda pt:len(pt)>0,txt.split('\n'))
    try:
        #title = pts[0]
        authors = pts[1]
        affs = ''.join(pts[2:])
        return [authors,affs]
        #content = pts[-1]
    except IndexError:
        print 'index'
        return ['NOT_USE_TEMPLATE','NOT_USE_TEMPLATE']


def parse_docx(filename):
    document = Document(filename)
    txt = []
    title = None
    authors = None
    affs = []

    for para in  document.paragraphs:
        if len(para.text)<=0:
            continue
        if para.text.startswith('Abstract'):
            break
        if title is None:
            title= para.text
        elif authors is None:
            authors = para.text
        else:
            affs.append(para.text)
    return [authors,''.join(affs)]

if __name__ == "__main__":
    #txt = doc_to_text_antiword("/tmp/somedocx.docx")
    #title,authors,affs,content = parse_user_name_aff(txt)
    authors ,affs = parse_docx('/tmp/143.docx')
    print 'authors:',authors
    print 'affs:',affs
